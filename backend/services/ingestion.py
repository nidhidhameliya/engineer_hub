"""Text extraction service for all supported file types."""
import base64
import io
import json
from pathlib import Path
from typing import Optional

import structlog
from openai import AsyncOpenAI

from config import get_settings

logger = structlog.get_logger()
settings = get_settings()


async def extract_text(file_path: Path, mime_type: Optional[str] = None) -> str:
    """Route to appropriate extractor based on file extension."""
    suffix = file_path.suffix.lower()

    extractors = {
        ".pdf": _extract_pdf,
        ".docx": _extract_docx,
        ".doc": _extract_docx,
        ".txt": _extract_text_plain,
        ".md": _extract_text_plain,
        ".markdown": _extract_text_plain,
        ".json": _extract_json,
        ".csv": _extract_csv,
        ".png": _extract_image_vision,
        ".jpg": _extract_image_vision,
        ".jpeg": _extract_image_vision,
    }

    extractor = extractors.get(suffix)
    if extractor is None:
        logger.warning("No extractor for file type", suffix=suffix)
        return _extract_text_plain(file_path)

    try:
        return await extractor(file_path)
    except Exception as e:
        logger.error("Extraction failed", file=str(file_path), error=str(e))
        raise


def _extract_pdf_sync(file_path: Path) -> str:
    """Synchronous PDF extraction using pypdf (pure Python, works on all Python versions)."""
    from pypdf import PdfReader
    reader = PdfReader(str(file_path))
    pages = []
    for page_num, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        if text.strip():
            pages.append(f"[Page {page_num + 1}]\n{text}")
    return "\n\n".join(pages)


async def _extract_pdf(file_path: Path) -> str:
    import asyncio
    loop = asyncio.get_event_loop()
    return await loop.run_in_executor(None, _extract_pdf_sync, file_path)


def _extract_docx_sync(file_path: Path) -> str:
    from docx import Document
    doc = Document(str(file_path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)
    return "\n\n".join(paragraphs)


async def _extract_docx(file_path: Path) -> str:
    import asyncio
    loop = asyncio.get_event_loop()
    return await loop.run_in_executor(None, _extract_docx_sync, file_path)


async def _extract_text_plain(file_path: Path) -> str:
    try:
        return file_path.read_text(encoding="utf-8", errors="replace")
    except Exception:
        return file_path.read_text(encoding="latin-1", errors="replace")


async def _extract_json(file_path: Path) -> str:
    try:
        data = json.loads(file_path.read_text(encoding="utf-8"))
        return json.dumps(data, indent=2)
    except Exception:
        return file_path.read_text(encoding="utf-8", errors="replace")


async def _extract_csv(file_path: Path) -> str:
    import csv
    lines = []
    try:
        with open(file_path, newline="", encoding="utf-8", errors="replace") as f:
            reader = csv.reader(f)
            for row in reader:
                lines.append(", ".join(row))
        return "\n".join(lines)
    except Exception:
        return file_path.read_text(encoding="utf-8", errors="replace")


async def _extract_image_vision(file_path: Path) -> str:
    """Use GPT-4o Vision to extract text and structure from architecture images."""
    client = AsyncOpenAI(api_key=settings.openai_api_key)

    with open(file_path, "rb") as f:
        image_data = base64.b64encode(f.read()).decode("utf-8")

    ext = file_path.suffix.lower().lstrip(".")
    mime_map = {"jpg": "jpeg", "jpeg": "jpeg", "png": "png"}
    mime = mime_map.get(ext, "png")

    prompt = (
        "You are an engineering documentation assistant. "
        "Analyze this architecture or system diagram carefully. "
        "Extract and describe:\n"
        "1. All service/component names\n"
        "2. Relationships and data flows between services\n"
        "3. Databases and storage systems\n"
        "4. External integrations and APIs\n"
        "5. Infrastructure components (load balancers, queues, etc.)\n"
        "6. Any text labels, annotations, or notes\n\n"
        "Format as structured text that can be searched later. "
        "Be thorough and precise."
    )

    try:
        response = await client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:image/{mime};base64,{image_data}",
                                "detail": "high",
                            },
                        },
                        {"type": "text", "text": prompt},
                    ],
                }
            ],
            max_tokens=2000,
        )
        extracted = response.choices[0].message.content or ""
        return f"[Architecture Diagram: {file_path.name}]\n\n{extracted}"
    except Exception as e:
        logger.error("Vision extraction failed", error=str(e))
        return f"[Image: {file_path.name}] - Vision extraction failed: {str(e)}"
